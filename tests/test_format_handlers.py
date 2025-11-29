"""Tests for format handlers with vocabulary rendering."""

import pytest
from pathlib import Path
from io import BytesIO

from anchor_text.formatting.ir import (
    FormattedDocument,
    TextBlock,
    TextStyle,
    VocabularyMetadata,
    LexicalMap,
    WordEntry,
    MorphemeInfo,
    MorphemeFamily,
    DecoderTrap,
    TrapOption,
)
from anchor_text.formats.pdf_handler import (
    PDFHandler,
    DIFFICULTY_COLORS,
    TRAP_BG_COLOR,
    VOCAB_HEADER_COLOR,
)
from anchor_text.formats.epub_handler import EPUBHandler
from anchor_text.formats.docx_handler import DOCXHandler


class TestPDFHandlerVocabulary:
    """Tests for PDF handler vocabulary rendering."""

    @pytest.fixture
    def handler(self):
        return PDFHandler()

    @pytest.fixture
    def sample_document(self):
        """Create a sample document with vocabulary metadata."""
        doc = FormattedDocument()

        # Add some text blocks
        block1 = TextBlock()
        block1.append("The ", TextStyle.NONE)
        block1.append("scientists", TextStyle.BOLD)
        block1.append(" hypothesized about the results.", TextStyle.NONE)
        doc.blocks.append(block1)

        # Add a decoder trap block
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What did the scientists do?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True
        doc.blocks.append(trap_block)

        # Add vocabulary metadata
        doc.vocabulary = VocabularyMetadata()
        lexical_map = LexicalMap()

        # Add some words to the lexical map
        word1 = WordEntry(
            word="hypothesized",
            root="hypo",
            syllables=["hy", "poth", "e", "sized"],
            difficulty_score=7,
            morphemes=[
                MorphemeInfo(text="hypo", meaning="under", origin="Greek", morpheme_type="prefix"),
                MorphemeInfo(text="thes", meaning="put", origin="Greek", morpheme_type="root"),
            ]
        )
        word2 = WordEntry(
            word="scientists",
            root="sci",
            syllables=["sci", "en", "tists"],
            difficulty_score=5,
            morphemes=[
                MorphemeInfo(text="sci", meaning="knowledge", origin="Latin", morpheme_type="root"),
            ]
        )
        word3 = WordEntry(
            word="results",
            root="sult",
            syllables=["re", "sults"],
            difficulty_score=3,
        )

        lexical_map.add_word(word1)
        lexical_map.add_word(word2)
        lexical_map.add_word(word3)
        doc.vocabulary.lexical_map = lexical_map

        return doc

    def test_creates_styles(self, handler):
        """Test that all required styles are created."""
        styles = handler._create_styles()

        assert "body" in styles
        assert "trap" in styles
        assert "vocab_header" in styles
        assert "vocab_subheader" in styles
        assert "vocab_word" in styles
        assert "family_header" in styles

    def test_block_to_html_escapes_special_chars(self, handler):
        """Test HTML special characters are escaped."""
        block = TextBlock()
        block.append("<script>alert('xss')</script>", TextStyle.NONE)

        html = handler._block_to_html(block)

        assert "&lt;" in html
        assert "&gt;" in html
        assert "<script>" not in html

    def test_block_to_html_formats_bold_italic(self, handler):
        """Test bold and italic formatting."""
        block = TextBlock()
        block.append("normal ", TextStyle.NONE)
        block.append("bold ", TextStyle.BOLD)
        block.append("italic ", TextStyle.ITALIC)
        block.append("both", TextStyle.BOLD | TextStyle.ITALIC)

        html = handler._block_to_html(block)

        assert "normal" in html
        assert "<b>bold </b>" in html
        assert "<i>italic </i>" in html
        assert "<b><i>both</i></b>" in html

    def test_format_morphemes(self, handler):
        """Test morpheme formatting."""
        entry = WordEntry(
            word="unpredictable",
            root="dict",
            morphemes=[
                MorphemeInfo(text="un", meaning="not", morpheme_type="prefix"),
                MorphemeInfo(text="pre", meaning="before", morpheme_type="prefix"),
                MorphemeInfo(text="dict", meaning="say", morpheme_type="root"),
                MorphemeInfo(text="able", meaning="capable of", morpheme_type="suffix"),
            ]
        )

        result = handler._format_morphemes(entry)

        assert "un (not)" in result
        assert "pre (before)" in result
        assert "dict (say)" in result
        assert " + " in result

    def test_format_morphemes_empty(self, handler):
        """Test morpheme formatting with no morphemes."""
        entry = WordEntry(word="test", root="test", morphemes=[])
        result = handler._format_morphemes(entry)
        assert result == ""

    def test_render_vocabulary_section(self, handler, sample_document):
        """Test vocabulary section rendering creates elements."""
        styles = handler._create_styles()
        elements = handler._render_vocabulary_section(
            sample_document.vocabulary, styles
        )

        # Should have header and tier elements
        assert len(elements) > 0

    def test_render_trap_block(self, handler):
        """Test decoder trap block rendering."""
        styles = handler._create_styles()

        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What is the answer?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True

        elements = handler._render_trap_block(trap_block, styles)

        # Should have spacer, table, spacer
        assert len(elements) == 3

    def test_render_word_families(self, handler):
        """Test word families rendering."""
        styles = handler._create_styles()

        # Create sample families
        families = [
            MorphemeFamily(
                root=MorphemeInfo(text="dict", meaning="say", origin="Latin"),
                words=[
                    WordEntry(word="predict", syllables=["pre", "dict"]),
                    WordEntry(word="dictate", syllables=["dic", "tate"]),
                ]
            )
        ]

        elements = handler._render_word_families(families, styles)

        # Should have header and family content
        assert len(elements) > 0

    def test_write_creates_pdf_with_vocabulary(self, handler, sample_document, tmp_path):
        """Test PDF writing includes vocabulary section."""
        output_path = tmp_path / "output.pdf"

        handler.write(sample_document, output_path)

        assert output_path.exists()
        # PDF should have non-trivial size due to vocabulary
        assert output_path.stat().st_size > 1000

    def test_write_creates_pdf_without_vocabulary(self, handler, tmp_path):
        """Test PDF writing works without vocabulary metadata."""
        doc = FormattedDocument()
        block = TextBlock()
        block.append("Simple text without vocabulary.", TextStyle.NONE)
        doc.blocks.append(block)

        output_path = tmp_path / "simple.pdf"
        handler.write(doc, output_path)

        assert output_path.exists()


class TestPDFDifficultyColors:
    """Tests for difficulty color constants."""

    def test_all_tiers_have_colors(self):
        """Test that all difficulty tiers have defined colors."""
        assert "easy" in DIFFICULTY_COLORS
        assert "medium" in DIFFICULTY_COLORS
        assert "challenging" in DIFFICULTY_COLORS

    def test_trap_and_header_colors_exist(self):
        """Test special colors are defined."""
        assert TRAP_BG_COLOR is not None
        assert VOCAB_HEADER_COLOR is not None


class TestEPUBHandlerDecoderTraps:
    """Tests for EPUB handler decoder trap rendering."""

    @pytest.fixture
    def handler(self):
        return EPUBHandler()

    @pytest.fixture
    def sample_document_with_vocab(self):
        """Create a sample document with vocabulary metadata."""
        doc = FormattedDocument()

        # Add some text blocks
        block1 = TextBlock()
        block1.append("The scientists hypothesized about the results.", TextStyle.NONE)
        doc.blocks.append(block1)

        # Add a decoder trap block
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What did the scientists do?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True
        doc.blocks.append(trap_block)

        # Add vocabulary metadata
        doc.vocabulary = VocabularyMetadata()
        lexical_map = LexicalMap()

        word1 = WordEntry(
            word="hypothesized",
            root="hypo",
            syllables=["hy", "poth", "e", "sized"],
            difficulty_score=7,
            morphemes=[
                MorphemeInfo(text="hypo", meaning="under", origin="Greek", morpheme_type="prefix"),
            ]
        )
        lexical_map.add_word(word1)
        doc.vocabulary.lexical_map = lexical_map

        return doc

    def test_document_to_html_marks_traps(self, handler):
        """Test decoder traps get special CSS class."""
        doc = FormattedDocument()

        # Normal block
        block1 = TextBlock()
        block1.append("Normal paragraph.", TextStyle.NONE)
        doc.blocks.append(block1)

        # Trap block
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: Question here]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True
        doc.blocks.append(trap_block)

        html = handler._document_to_html(doc)

        assert 'class="decoder-trap"' in html
        assert "Decoder Check" in html

    def test_render_trap_block_html_creates_details(self, handler):
        """Test decoder trap renders with details element."""
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What word means happy?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True

        html = handler._render_trap_block_html(trap_block)

        assert "<details>" in html
        assert "<summary>" in html
        assert "What word means happy?" in html

    def test_render_trap_block_html_escapes_special_chars(self, handler):
        """Test HTML special characters are escaped in traps."""
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What is < or > sign?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True

        html = handler._render_trap_block_html(trap_block)

        assert "&lt;" in html
        assert "&gt;" in html

    def test_escape_html(self, handler):
        """Test HTML escaping utility."""
        assert handler._escape_html("<test>") == "&lt;test&gt;"
        assert handler._escape_html("a & b") == "a &amp; b"
        assert handler._escape_html("normal") == "normal"

    def test_render_vocabulary_section_html(self, handler, sample_document_with_vocab):
        """Test vocabulary section HTML generation."""
        lexical_map = sample_document_with_vocab.vocabulary.lexical_map
        html = handler._render_vocabulary_section_html(lexical_map)

        assert 'class="vocab-section"' in html
        assert "Vocabulary Preview" in html
        # Should have the challenging word
        assert "hy·poth·e·sized" in html or "hypothesized" in html

    def test_format_morphemes_html(self, handler):
        """Test morpheme HTML formatting."""
        entry = WordEntry(
            word="predict",
            root="dict",
            morphemes=[
                MorphemeInfo(text="pre", meaning="before"),
                MorphemeInfo(text="dict", meaning="say"),
            ]
        )

        html = handler._format_morphemes_html(entry)

        assert "pre (before)" in html
        assert "dict (say)" in html
        assert " + " in html

    def test_render_word_families_html(self, handler):
        """Test word families HTML generation."""
        families = [
            MorphemeFamily(
                root=MorphemeInfo(text="dict", meaning="say", origin="Latin"),
                words=[
                    WordEntry(word="predict", syllables=["pre", "dict"]),
                    WordEntry(word="dictate", syllables=["dic", "tate"]),
                ]
            )
        ]

        html = handler._render_word_families_html(families)

        assert 'class="word-families"' in html
        assert "DICT" in html
        assert "say" in html
        assert "Latin" in html

    def test_document_to_html_includes_vocabulary(self, handler, sample_document_with_vocab):
        """Test full document HTML includes vocabulary section."""
        html = handler._document_to_html(sample_document_with_vocab)

        assert 'class="vocab-section"' in html
        assert 'class="decoder-trap"' in html
        assert "<details>" in html

    def test_write_creates_epub_with_vocabulary(self, handler, sample_document_with_vocab, tmp_path):
        """Test EPUB writing includes vocabulary content."""
        output_path = tmp_path / "output.epub"

        handler.write(sample_document_with_vocab, output_path)

        assert output_path.exists()
        assert output_path.stat().st_size > 1000


class TestDOCXHandlerVocabulary:
    """Tests for DOCX handler vocabulary rendering."""

    @pytest.fixture
    def handler(self):
        return DOCXHandler()

    @pytest.fixture
    def sample_document_with_vocab(self):
        """Create a sample document with vocabulary metadata."""
        doc = FormattedDocument()

        # Add some text blocks
        block1 = TextBlock()
        block1.append("The scientists hypothesized about the results.", TextStyle.NONE)
        doc.blocks.append(block1)

        # Add a decoder trap block
        trap_block = TextBlock()
        trap_block.append("[Decoder Check: What did the scientists do?]", TextStyle.ITALIC)
        trap_block.is_decoder_trap = True
        doc.blocks.append(trap_block)

        # Add vocabulary metadata
        doc.vocabulary = VocabularyMetadata()
        lexical_map = LexicalMap()

        word1 = WordEntry(
            word="hypothesized",
            root="hypo",
            syllables=["hy", "poth", "e", "sized"],
            difficulty_score=7,
            morphemes=[
                MorphemeInfo(text="hypo", meaning="under", origin="Greek", morpheme_type="prefix"),
            ]
        )
        word2 = WordEntry(
            word="scientists",
            root="sci",
            syllables=["sci", "en", "tists"],
            difficulty_score=5,
        )
        lexical_map.add_word(word1)
        lexical_map.add_word(word2)
        doc.vocabulary.lexical_map = lexical_map

        return doc

    def test_write_creates_docx(self, handler, tmp_path):
        """Test basic DOCX writing."""
        doc = FormattedDocument()
        block = TextBlock()
        block.append("Test content.", TextStyle.NONE)
        doc.blocks.append(block)

        output_path = tmp_path / "test.docx"
        handler.write(doc, output_path)

        assert output_path.exists()

    def test_write_creates_docx_with_vocabulary(self, handler, sample_document_with_vocab, tmp_path):
        """Test DOCX writing with vocabulary content."""
        output_path = tmp_path / "vocab_test.docx"

        handler.write(sample_document_with_vocab, output_path)

        assert output_path.exists()
        # DOCX should have non-trivial size due to vocabulary tables
        assert output_path.stat().st_size > 5000

    def test_format_morphemes(self, handler):
        """Test morpheme formatting."""
        entry = WordEntry(
            word="unpredictable",
            root="dict",
            morphemes=[
                MorphemeInfo(text="un", meaning="not"),
                MorphemeInfo(text="pre", meaning="before"),
                MorphemeInfo(text="dict", meaning="say"),
            ]
        )

        result = handler._format_morphemes(entry)

        assert "un (not)" in result
        assert "pre (before)" in result
        assert "dict (say)" in result
        assert " + " in result

    def test_format_morphemes_empty(self, handler):
        """Test morpheme formatting with no morphemes."""
        entry = WordEntry(word="test", root="test", morphemes=[])
        result = handler._format_morphemes(entry)
        assert result == ""

    def test_docx_contains_decoder_trap(self, handler, sample_document_with_vocab, tmp_path):
        """Test that DOCX contains decoder trap content."""
        output_path = tmp_path / "trap_test.docx"
        handler.write(sample_document_with_vocab, output_path)

        # Read back and check content
        from docx import Document as DocxDocument
        read_doc = DocxDocument(output_path)

        # Find the decoder check text
        found_trap = False
        for para in read_doc.paragraphs:
            if "Decoder Check" in para.text:
                found_trap = True
                break

        # Also check tables for trap content
        for table in read_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Decoder Check" in cell.text:
                        found_trap = True

        assert found_trap

    def test_docx_contains_vocabulary_header(self, handler, sample_document_with_vocab, tmp_path):
        """Test that DOCX contains vocabulary section header."""
        output_path = tmp_path / "header_test.docx"
        handler.write(sample_document_with_vocab, output_path)

        from docx import Document as DocxDocument
        read_doc = DocxDocument(output_path)

        found_header = False
        for para in read_doc.paragraphs:
            if "VOCABULARY PREVIEW" in para.text:
                found_header = True
                break

        assert found_header
