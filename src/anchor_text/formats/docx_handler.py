"""Microsoft Word (.docx) file handler."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from anchor_text.formats.base import FormatHandler
from anchor_text.formatting.ir import (
    FormattedDocument,
    ImageRef,
    TextBlock,
    VocabularyMetadata,
    WordEntry,
    MorphemeFamily,
)


# Color definitions for difficulty tiers
DIFFICULTY_COLORS = {
    "easy": RGBColor(232, 245, 233),      # Light green
    "medium": RGBColor(255, 243, 224),    # Light orange
    "challenging": RGBColor(255, 235, 238),  # Light red
}
TRAP_BG_COLOR = RGBColor(227, 242, 253)  # Light blue
VOCAB_HEADER_COLOR = RGBColor(25, 118, 210)  # Blue


class DOCXHandler(FormatHandler):
    """Handler for Microsoft Word (.docx) files.

    Uses python-docx for reading and writing with full support for
    bold and italic text formatting at the run level.
    """

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def read(self, path: Path) -> str:
        """Extract plain text from DOCX file."""
        doc = Document(path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def write(self, document: FormattedDocument, path: Path) -> None:
        """Write formatted document to DOCX file.

        Creates a new Word document with run-level formatting for
        bold and italic text, vocabulary tables, and styled decoder traps.
        """
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Add vocabulary preview section at the beginning
        if document.vocabulary and document.vocabulary.lexical_map:
            self._add_vocabulary_section(doc, document.vocabulary)
            # Add horizontal line separator
            self._add_horizontal_line(doc)

        for block in document.blocks:
            if block.is_decoder_trap:
                # Render decoder trap with styled box
                self._add_decoder_trap(doc, block)
            else:
                para = doc.add_paragraph()
                for run_data in block.runs:
                    run = para.add_run(run_data.text)
                    run.bold = run_data.bold
                    run.italic = run_data.italic

        # Add word families section at the end
        if document.vocabulary and document.vocabulary.lexical_map:
            families = document.vocabulary.lexical_map.get_root_families()
            if families:
                self._add_horizontal_line(doc)
                self._add_word_families(doc, families)

        # Insert images at their positions
        # Note: This is simplified - images go at the end
        # More sophisticated positioning would require tracking
        # exact positions during text extraction
        for image in document.images:
            doc.add_paragraph()  # Add spacing
            para = doc.add_paragraph()
            run = para.add_run()

            # Save image to temp file and add to document
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(
                suffix=f".{image.format}", delete=False
            ) as f:
                f.write(image.data)
                temp_path = f.name

            try:
                run.add_picture(temp_path)
            finally:
                os.unlink(temp_path)

        doc.save(path)

    def _add_vocabulary_section(
        self, doc: Document, vocab: VocabularyMetadata
    ) -> None:
        """Add vocabulary preview section to the document."""
        # Header
        header = doc.add_paragraph()
        header_run = header.add_run("VOCABULARY PREVIEW")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.color.rgb = VOCAB_HEADER_COLOR

        lexical_map = vocab.lexical_map
        if not lexical_map:
            return

        # Render each difficulty tier
        for tier_name, tier_label in [
            ("challenging", "Challenging Words"),
            ("medium", "Medium Words"),
            ("easy", "Familiar Words"),
        ]:
            tier_words = lexical_map.difficulty_tiers.get(tier_name, [])
            if not tier_words:
                continue

            # Tier header
            tier_header = doc.add_paragraph()
            tier_run = tier_header.add_run(f"{tier_label} ({len(tier_words)} words)")
            tier_run.bold = True
            tier_run.font.size = Pt(11)

            # Create table for this tier
            display_words = tier_words[:10]
            if display_words:
                table = doc.add_table(rows=len(display_words), cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, word_key in enumerate(display_words):
                    entry = lexical_map.words.get(word_key)
                    if entry:
                        # Word with syllables
                        word_cell = table.rows[i].cells[0]
                        word_para = word_cell.paragraphs[0]
                        word_run = word_para.add_run(entry.syllable_text)
                        word_run.bold = True

                        # Morpheme breakdown
                        morpheme_cell = table.rows[i].cells[1]
                        morpheme_para = morpheme_cell.paragraphs[0]
                        morpheme_text = self._format_morphemes(entry)
                        morpheme_run = morpheme_para.add_run(morpheme_text)
                        morpheme_run.italic = True

                        # Set cell background color
                        tier_color = DIFFICULTY_COLORS[tier_name]
                        self._set_cell_shading(word_cell, tier_color)
                        self._set_cell_shading(morpheme_cell, tier_color)

            if len(tier_words) > 10:
                more_para = doc.add_paragraph()
                more_run = more_para.add_run(
                    f"... and {len(tier_words) - 10} more {tier_name} words"
                )
                more_run.italic = True

    def _format_morphemes(self, entry: WordEntry) -> str:
        """Format morpheme breakdown for display."""
        if not entry.morphemes:
            return ""
        parts = []
        for m in entry.morphemes:
            if m.meaning:
                parts.append(f"{m.text} ({m.meaning})")
            else:
                parts.append(m.text)
        return " + ".join(parts)

    def _add_decoder_trap(self, doc: Document, block: TextBlock) -> None:
        """Add a styled decoder trap section."""
        # Create a single-cell table to act as a styled box
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        self._set_cell_shading(cell, TRAP_BG_COLOR)

        # Add the trap content
        para = cell.paragraphs[0]
        for run_data in block.runs:
            run = para.add_run(run_data.text)
            run.bold = run_data.bold
            run.italic = run_data.italic
            run.font.color.rgb = VOCAB_HEADER_COLOR

        # Add border to table
        self._set_table_border(table, VOCAB_HEADER_COLOR)

        # Add spacing after
        doc.add_paragraph()

    def _add_word_families(
        self, doc: Document, families: list[MorphemeFamily]
    ) -> None:
        """Add word families section."""
        # Header
        header = doc.add_paragraph()
        header_run = header.add_run("WORD FAMILIES")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.color.rgb = VOCAB_HEADER_COLOR

        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("Words grouped by their root morpheme")
        sub_run.italic = True

        for family in families[:8]:
            root = family.root
            root_text = root.text.upper()
            if root.meaning:
                root_text += f" ({root.meaning})"
            if root.origin:
                root_text += f" - {root.origin}"

            # Family header
            family_para = doc.add_paragraph()
            family_run = family_para.add_run(root_text)
            family_run.bold = True
            family_run.font.color.rgb = VOCAB_HEADER_COLOR

            # Word list
            word_list = ", ".join(w.syllable_text for w in family.words[:6])
            if len(family.words) > 6:
                word_list += f", ... (+{len(family.words) - 6} more)"

            words_para = doc.add_paragraph()
            words_para.paragraph_format.left_indent = Inches(0.5)
            words_para.add_run(word_list)

    def _add_horizontal_line(self, doc: Document) -> None:
        """Add a horizontal line separator."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        # Create horizontal line using paragraph border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_cell_shading(self, cell, color: RGBColor) -> None:
        """Set background color for a table cell."""
        shading_elm = OxmlElement("w:shd")
        # RGBColor is a tuple-like object, use indexing to get RGB values
        shading_elm.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_table_border(self, table, color: RGBColor) -> None:
        """Set border for a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")

        # RGBColor is a tuple-like object, use indexing to get RGB values
        color_str = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color_str)
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def extract_images(self, path: Path) -> list[ImageRef]:
        """Extract images from DOCX file."""
        import zipfile

        images: list[ImageRef] = []
        position = 0

        # DOCX is a ZIP file - extract images from media folder
        with zipfile.ZipFile(path, "r") as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    data = zf.read(name)
                    # Determine format from filename
                    ext = Path(name).suffix.lower().lstrip(".")
                    if ext in ("png", "jpg", "jpeg", "gif", "bmp"):
                        if ext == "jpeg":
                            ext = "jpg"
                        images.append(
                            ImageRef(
                                data=data,
                                format=ext,
                                position=position,
                            )
                        )
                        position += 1

        return images
